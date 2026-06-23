"""生成 Word 解题报告"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

DIR    = Path(__file__).parent
OUTPUT = DIR / "解题报告_强化学习动态因子.docx"


def font(run, name="微软雅黑", size=11, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size  = Pt(size)
    run.font.bold  = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for r in p.runs:
        r.font.name = "微软雅黑"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def para(doc, text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    if indent:
        p.paragraph_format.left_indent = Cm(0.6)
    r = p.add_run(text)
    font(r)
    return p


def code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "F2F2F2")
    pPr.append(shd)
    r = p.add_run(text)
    r.font.name = "Courier New"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(9)


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style     = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in c.paragraphs[0].runs:
            font(r, bold=True, size=10)
        pr = c._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  "D9E1F2")
        pr.append(shd)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri + 1].cells[ci]
            c.text = str(val)
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in c.paragraphs[0].runs:
                font(r, size=10)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def pic(doc, path, width_cm=14, caption=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    if caption:
        cp = doc.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in cp.runs:
            font(r, size=9, color=(100, 100, 100))


# ── 正文 ──────────────────────────────────────────────────────

def build():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin    = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin   = Cm(3.0)
    sec.right_margin  = Cm(2.5)

    # 标题
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("题目三：强化学习动态因子选择")
    font(r, size=16, bold=True, color=(31, 73, 125))
    doc.add_paragraph()

    # ── 一、思路 ──────────────────────────────────
    heading(doc, "一、思路", 1)

    para(doc,
        "因子有效性随市场状态变化，动量因子在趋势行情有效，"
        "反转因子在震荡行情有效，固定权重无法适应这种切换。")

    para(doc,
        "直接用 XGBoost 预测各因子下期 IC、按 IC 大小分配权重是可行的，"
        "但这里用强化学习（FQI）做了一层粗分类——先判断当前适合哪类配方，"
        "再用 XGBoost 在配方内微调权重。两层分工：")

    code(doc,
        "第一层（FQI）：观察市场状态 → 选配方（动量型 / 反转型 / 等权...）\n"
        "第二层（XGBoost）：选好配方后 → 预测配方内各因子下期IC → 精确权重")

    para(doc,
        "选 FQI 而不是 DQN 等深度RL的原因：月频数据约70个训练期，"
        "样本太少，XGBoost 比神经网络更稳；FQI 本质上就是反复做 XGBoost 回归，"
        "只是每轮的训练标签在被上一轮结果修正。")

    # ── 二、数据生成 ──────────────────────────────
    heading(doc, "二、模拟数据生成（gen_data.py）", 1)

    para(doc, "300只股票 × 1500个交易日，10个候选因子，2种市场状态（Markov链切换）。")

    para(doc, "核心设计：因子有效性在两种状态下方向相反，模拟真实的 Regime 切换。")

    table(doc,
        ["", "F0 F1 F2（动量类）", "F3 F4 F6（反转类）", "其余因子"],
        [
            ["趋势行情（Regime 0）", "IC ≈ +0.07~+0.09", "IC ≈ −0.04~−0.05", "混合"],
            ["震荡行情（Regime 1）", "IC ≈ −0.04~−0.06", "IC ≈ +0.07~+0.09", "混合"],
        ],
        widths=[3.5, 4.5, 4.5, 3.0],
    )
    doc.add_paragraph()

    para(doc,
        "收益 = 市场公共涨跌 + Σ(因子 × 当前状态载荷) + 个股噪声(σ=0.03)。"
        "因子间有截面相关结构（公共成分0.4）+ 时序自相关（AR系数0.3），"
        "接近真实数据特征。")

    # ── 三、环境设计 ──────────────────────────────
    heading(doc, "三、环境设计（rl_env.py）", 1)

    heading(doc, "状态（13维）", 2)
    table(doc,
        ["维度", "内容", "作用"],
        [
            ["1~10", "各因子过去20日滚动RankIC", "趋势行情时动量IC走高，震荡时反转IC走高，间接反映市场状态"],
            ["11",   "市场截面收益率波动率",      "高波动 = 震荡分化"],
            ["12",   "市场近5日平均收益",          "短期趋势方向"],
            ["13",   "因子间平均相关系数",          "因子分散度，高相关时等权效果差"],
        ],
        widths=[1.5, 5.0, 8.0],
    )
    doc.add_paragraph()

    heading(doc, "动作（5种配方）", 2)
    table(doc,
        ["动作", "配方", "权重"],
        [
            ["0", "等权",     "10个因子各1/10"],
            ["1", "动量型",   "F0/F1/F2/F5 各1/4"],
            ["2", "反转型",   "F3/F4/F6 各1/3"],
            ["3", "波动型",   "F7/F8/F9 各1/3"],
            ["4", "IC自适应", "按滚动IC绝对值加权"],
        ],
        widths=[1.5, 3.0, 10.0],
    )
    doc.add_paragraph()

    heading(doc, "奖励", 2)
    para(doc,
        "选定配方后持仓20天（月频），奖励 = 这20天多空组合的平均RankIC。"
        "用RankIC而不是实际收益，是为了剥离市场Beta，只看因子截面预测能力。")

    # ── 四、算法 ──────────────────────────────────
    heading(doc, "四、FQI 关键伪代码", 1)

    code(doc,
        "# 初始化：用随机策略收集经验 (状态, 选了哪个配方, 赚了多少, 下期状态)\n"
        "D = collect_with_random_policy(env)\n"
        "Q = {每个配方: 全零初始化}\n\n"
        "for k in range(20):\n"
        "    # Bellman 目标：当期收益 + 折扣 × 下期最优Q值\n"
        "    y = reward + 0.95 * max_a(Q[a](next_state))\n\n"
        "    # 按配方分组，各自训练一个 XGBoost\n"
        "    for 配方 in [0,1,2,3,4]:\n"
        "        Q[配方] = XGBRegressor().fit(states, y)  # 普通回归\n\n"
        "    # 用更新后的策略再收集数据，扩充经验池\n"
        "    D += collect_with_epsilon_greedy(env, Q)\n\n"
        "# 决策：给定状态，取Q值最大的配方\n"
        "action = argmax([Q[a].predict(state) for a in range(5)])")

    para(doc,
        "每一步都是普通 XGBoost.fit()，不同的只是训练标签 y 在被上一轮预测不断修正，"
        "直到 Q 值稳定。")

    # ── 五、结果 ──────────────────────────────────
    heading(doc, "五、结果", 1)

    heading(doc, "单层 FQI vs 各基准（测试集22期）", 2)
    table(doc,
        ["策略", "平均RankIC", "ICIR", "IC>0%", "累积IC"],
        [
            ["等权基准",    "0.616", "17.49", "100%", "13.55"],
            ["IC自适应",    "0.615", "17.56", "100%", "13.53"],
            ["FQI智能体",   "0.587",  "9.30", "100%", "12.90"],
            ["动量型（固定）","0.562", "6.72", "100%", "12.36"],
            ["随机策略",    "0.597", "12.00", "100%", "13.14"],
        ],
        widths=[4.0, 3.0, 2.5, 2.5, 2.5],
    )
    doc.add_paragraph()

    if (DIR / "rl_results.png").exists():
        pic(doc, DIR / "rl_results.png", width_cm=15,
            caption="图1  各策略累积RankIC曲线 / 每期IC分布 / FQI动作分布")

    heading(doc, "两层系统（FQI选配方 + XGBoost微调权重）", 2)
    table(doc,
        ["策略", "平均RankIC", "ICIR", "累积IC"],
        [
            ["等权基准",  "0.616", "17.49", "13.55"],
            ["FQI单层",   "0.587",  "9.30", "12.90"],
            ["两层系统",  "0.589",  "9.36", "12.97"],
        ],
        widths=[4.5, 3.5, 3.0, 3.5],
    )
    doc.add_paragraph()

    if (DIR / "two_layer_results.png").exists():
        pic(doc, DIR / "two_layer_results.png", width_cm=15,
            caption="图2  两层系统 vs 单层FQI vs 等权基准")

    if (DIR / "feature_importance.png").exists():
        pic(doc, DIR / "feature_importance.png", width_cm=12,
            caption="图3  Q函数特征重要性（各动作XGBoost均值）")

    # ── 六、分析 ──────────────────────────────────
    heading(doc, "六、结果分析与局限性", 1)

    heading(doc, "为什么等权基准最好", 2)
    para(doc,
        "模拟数据里 Regime 切换频繁，两种行情下各有一批因子有效，"
        "等权相当于「都压一点」，在有限样本下反而比择时更稳。"
        "FQI 的样本只有约70个月频训练期，Q函数估计不稳定，"
        "Bellman残差从0.59持续升到2.87，说明迭代在发散而非收敛。")

    heading(doc, "两层结构提升有限的原因", 2)
    para(doc,
        "第二层 WeightRefiner 也用同一批训练数据预测IC，"
        "样本瓶颈没有解决，所以提升很小（ICIR 9.30 → 9.36）。"
        "实际数据中因子数量多（几十到几百个），先粗筛再精调的分层设计更有意义。")

    heading(doc, "主要缺陷", 2)
    table(doc,
        ["问题", "说明"],
        [
            ["样本量不足", "月频约70个训练期，Q函数容易过拟合，Bellman残差发散"],
            ["动作空间是人工设计的", "5个预设配方融入了人类先验，若先验错误则上限受限"],
            ["未计交易成本", "月频换手成本约0.2%双边，实际多空收益会折损"],
            ["市值行业未中性化", "配方偏向可能隐含风格暴露"],
        ],
        widths=[4.0, 10.5],
    )
    doc.add_paragraph()

    heading(doc, "改进方向", 2)
    para(doc, "1. 用 CQL（保守Q学习）压制 Bellman 发散，专门为小样本离线RL设计；")
    para(doc, "2. 放弃离散配方，改为10维连续权重，用 XGBoost 直接预测各因子IC后归一化——"
             "这本质上就是题目里「IC自适应」的非线性升级版，也是最实用的方向；")
    para(doc, "3. 加入宏观特征（利率曲线、VIX等）作为 Regime 先行指标，提升状态识别能力。")

    doc.save(OUTPUT)
    print(f"报告已生成：{OUTPUT}")


if __name__ == "__main__":
    build()
